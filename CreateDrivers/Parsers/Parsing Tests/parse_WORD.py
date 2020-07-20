from docx import Document

out = open('output.txt', 'w')
# creates a Document obj
doc = Document('Self Critique Sp 3.docx')
# creates a table obj, note tables[] tracks the number of tables in the doc
table = doc.tables[0]
print("Number of tables:", len(doc.tables))
# these nested loops gets the text from each cell for each row
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            out.write(paragraph.text)
            out.write('\t')
    out.write('\n')

out.close()
